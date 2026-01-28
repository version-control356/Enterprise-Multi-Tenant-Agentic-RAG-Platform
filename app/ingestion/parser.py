import os
import io
import logging
import pandas as pd
import docx
from pypdf import PdfReader
from bs4 import BeautifulSoup
from typing import List, Optional
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    Language
)

logger = logging.getLogger(__name__)


class UniversalDocumentParser:
    """Extracts raw text from multiple file types and applies structure-aware chunking."""

    @staticmethod
    def _extract_pdf_ocr(reader: PdfReader, file_bytes: bytes) -> Optional[str]:
        """Attempt to extract text from scanned/image-only PDF pages using pytesseract OCR."""
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            logger.debug("pytesseract or PIL is not installed; skipping PDF OCR fallback.")
            return None

        ocr_pages = []
        try:
            for page_idx, page in enumerate(reader.pages):
                page_text_pieces = []

                if hasattr(page, "images") and page.images:
                    for img_file in page.images:
                        try:
                            pil_img = Image.open(io.BytesIO(img_file.data))
                            extracted = pytesseract.image_to_string(pil_img).strip()
                            if extracted:
                                page_text_pieces.append(extracted)
                        except Exception as img_err:
                            logger.debug("Error processing PDF page image: %s", img_err)


                if not page_text_pieces:
                    try:
                        import pypdfium2
                        pdf_doc = pypdfium2.PdfDocument(file_bytes)
                        if page_idx < len(pdf_doc):
                            page_render = pdf_doc[page_idx].render(scale=2.0)
                            pil_img = page_render.to_pil()
                            extracted = pytesseract.image_to_string(pil_img).strip()
                            if extracted:
                                page_text_pieces.append(extracted)
                    except Exception as render_err:
                        logger.debug("PDF rasterization OCR attempt failed: %s", render_err)

                if page_text_pieces:
                    combined_page_text = "\n".join(page_text_pieces)
                    ocr_pages.append(f"[Page {page_idx + 1} - OCR]\n{combined_page_text}")

            if ocr_pages:
                logger.info("✅ Extracted text from scanned PDF via pytesseract OCR.")
                return "\n\n".join(ocr_pages)
        except Exception as ocr_err:
            logger.warning("pytesseract OCR extraction encountered an error: %s", ocr_err)

        return None

    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> str:
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            text_pages = []
            for page_idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_pages.append(f"[Page {page_idx + 1}]\n{page_text.strip()}")

            if not text_pages or sum(len(p) for p in text_pages) < 20:
                ocr_result = UniversalDocumentParser._extract_pdf_ocr(reader, file_bytes)
                if ocr_result:
                    return ocr_result
                if not text_pages:
                    raise ValueError(
                        "The PDF document does not contain extractable text (it may be scanned/image-only). "
                        "Scanned PDF processing requires pytesseract and Tesseract OCR."
                    )

            return "\n\n".join(text_pages)

        elif ext == ".docx":
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
            return text

        elif ext in [".xlsx", ".xls"]:
            df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            full_text = []
            for sheet, df in df_dict.items():
                full_text.append(f"--- Sheet: {sheet} ---")
                full_text.append(df.to_string(index=False))
            return "\n".join(full_text)

        elif ext == ".csv":
            df = pd.read_csv(io.BytesIO(file_bytes))
            return df.to_string(index=False)

        elif ext in [".html", ".htm"]:
            soup = BeautifulSoup(file_bytes, "html.parser")
            return soup.get_text(separator="\n")

        elif ext in [".txt", ".md", ".json"]:
            return file_bytes.decode("utf-8", errors="ignore")

        elif ext in [".py", ".js", ".java", ".cpp"]:
            return file_bytes.decode("utf-8", errors="ignore")

        else:
            raise ValueError(f"Unsupported file format extension: {ext}")

    @classmethod
    def chunk_document(
        cls, 
        text: str, 
        filename: str, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 150
    ) -> List[str]:
        ext = os.path.splitext(filename)[1].lower()


        code_languages = {
            ".py": Language.PYTHON,
            ".js": Language.JS,
            ".java": Language.JAVA,
            ".cpp": Language.CPP,
        }

        if ext in code_languages:
            splitter = RecursiveCharacterTextSplitter.from_language(
                language=code_languages[ext],
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
        else:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )

        raw_chunks = splitter.split_text(text)
        return [f"[File: {filename}]\n{chunk}" for chunk in raw_chunks if chunk.strip()]